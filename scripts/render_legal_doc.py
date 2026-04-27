"""Render structured legal documents to DOCX + PDF.

Reusable for NDAs, role memos, advisor agreements, and similar text-only
legal instruments. Reads a Python module that exposes TITLE, BLOCKS (a list
of (kind, text) tuples), and OUT_BASENAME, and writes DOCX + PDF to the
caller-specified directory.

Block kinds:
- ("h1", text)        - document title (centered, bold, large)
- ("h2", text)        - section heading (bold, numbered if leading digit)
- ("p", text)         - body paragraph
- ("party", lines[])  - party block (each line on its own line)
- ("sig", lines[])    - signature block (lines + signature line)
- ("spacer", "")      - vertical space
"""

from __future__ import annotations
import argparse
import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib import colors


def render_docx(blocks, title, out_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, content in blocks:
        if kind == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(16)
            doc.add_paragraph()  # spacer
        elif kind == "h2":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(11)
        elif kind == "p":
            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "party":
            for line in content:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(0)
            doc.add_paragraph()
        elif kind == "sig":
            for line in content:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(0)
            doc.add_paragraph()
        elif kind == "spacer":
            doc.add_paragraph()
        else:
            raise ValueError(f"Unknown block kind: {kind}")

    doc.save(out_path)


def render_pdf(blocks, title, out_path: Path) -> None:
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        title=title,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=18,
        textColor=colors.black,
    )
    h2_style = ParagraphStyle(
        "Section",
        parent=styles["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    party_style = ParagraphStyle(
        "Party",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=13,
        spaceAfter=2,
    )

    story = []
    for kind, content in blocks:
        if kind == "h1":
            story.append(Paragraph(content, title_style))
        elif kind == "h2":
            story.append(Paragraph(content, h2_style))
        elif kind == "p":
            story.append(Paragraph(content, body_style))
        elif kind == "party":
            block_paras = [Paragraph(line, party_style) for line in content]
            story.append(KeepTogether(block_paras))
            story.append(Spacer(1, 8))
        elif kind == "sig":
            block_paras = [Paragraph(line, party_style) for line in content]
            story.append(KeepTogether(block_paras))
            story.append(Spacer(1, 14))
        elif kind == "spacer":
            story.append(Spacer(1, 10))
        else:
            raise ValueError(f"Unknown block kind: {kind}")

    doc.build(story)


def render_markdown(blocks, title, out_path: Path) -> None:
    """Plain-text/markdown source — editable and diffable."""
    lines = [f"# {title}", ""]
    for kind, content in blocks:
        if kind == "h1":
            continue  # title already written
        elif kind == "h2":
            lines.append(f"## {content}")
            lines.append("")
        elif kind == "p":
            lines.append(content)
            lines.append("")
        elif kind == "party":
            lines.extend(content)
            lines.append("")
        elif kind == "sig":
            lines.extend(content)
            lines.append("")
        elif kind == "spacer":
            lines.append("")
    out_path.write_text("\n".join(lines))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", help="Path to Python module defining TITLE, BLOCKS, OUT_BASENAME")
    parser.add_argument("--out-dir", default=".", help="Output directory")
    args = parser.parse_args()

    src = Path(args.source).resolve()
    spec = importlib.util.spec_from_file_location("legal_doc_src", src)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    base = mod.OUT_BASENAME

    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"
    md_path = out_dir / f"{base}.md"

    render_markdown(mod.BLOCKS, mod.TITLE, md_path)
    render_docx(mod.BLOCKS, mod.TITLE, docx_path)
    render_pdf(mod.BLOCKS, mod.TITLE, pdf_path)

    print(f"  wrote {md_path}")
    print(f"  wrote {docx_path}")
    print(f"  wrote {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
